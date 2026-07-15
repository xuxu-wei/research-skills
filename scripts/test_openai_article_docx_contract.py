#!/usr/bin/env python3
"""Build and structurally verify a synthetic scientific DOCX package."""

from __future__ import annotations

import argparse
import hashlib
import json
import shutil
import subprocess
import sys
import tempfile
import zipfile
from pathlib import Path

import yaml
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


REPO = Path(__file__).resolve().parents[1]
CONTRACT = REPO / "research-skills-openai/skills/article-orchestrator/references/article-docx-delivery-contract.md"
SECTIONS = {
    "Introduction": "The study addresses a defined scientific question and cites Table 1.",
    "Methods": "We used a prespecified synthetic analysis for this format test.",
    "Results": "The primary synthetic result is shown in Table 1 and Figure 1.",
    "Discussion": "The result demonstrates DOCX content and display integration only.",
}
TABLE = [["Group", "Estimate", "95% CI"], ["A", "1.20", "1.05 to 1.37"], ["B", "0.95", "0.82 to 1.10"]]


def sha256(path: Path) -> str:
    return "sha256:" + hashlib.sha256(path.read_bytes()).hexdigest()


def require(value: bool, label: str) -> None:
    if not value:
        raise AssertionError(label)


def make_figure(path: Path) -> None:
    image = Image.new("RGB", (1000, 560), "white")
    draw = ImageDraw.Draw(image)
    label_font = ImageFont.load_default(size=26)
    title_font = ImageFont.load_default(size=34)
    draw.line((100, 470, 900, 470), fill="black", width=4)
    draw.line((100, 470, 100, 80), fill="black", width=4)
    draw.rectangle((230, 250, 390, 470), fill="#5B8FF9")
    draw.rectangle((610, 315, 770, 470), fill="#61DDAA")
    draw.text((250, 490), "Group A", fill="black", font=label_font)
    draw.text((630, 490), "Group B", fill="black", font=label_font)
    draw.text((300, 25), "Synthetic effect estimates", fill="black", font=title_font)
    image.save(path, dpi=(150, 150))


def make_markdown(path: Path) -> None:
    lines = ["# Synthetic scientific manuscript", ""]
    for heading, text in SECTIONS.items():
        lines.extend((f"## {heading}", "", text, ""))
    lines.extend((
        "Table 1. Synthetic estimates.",
        "",
        "| Group | Estimate | 95% CI |",
        "|---|---:|---|",
        "| A | 1.20 | 1.05 to 1.37 |",
        "| B | 0.95 | 0.82 to 1.10 |",
        "",
        "Figure 1. Synthetic effect estimates for format verification.",
        "",
    ))
    path.write_text("\n".join(lines), encoding="utf-8", newline="\n")


def make_table_source(path: Path) -> None:
    path.write_text("\n".join(",".join(row) for row in TABLE) + "\n", encoding="utf-8", newline="\n")


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        from docx.oxml import OxmlElement
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    from docx.oxml import OxmlElement
    for side in ("top", "left", "bottom", "right"):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), "100")
        node.set(qn("w:type"), "dxa")


def build_docx(path: Path, figure: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = Inches(0.8)
    section.left_margin = section.right_margin = Inches(0.9)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    title = doc.add_paragraph("Synthetic scientific manuscript")
    title.style = styles["Title"]
    for heading, text in SECTIONS.items():
        doc.add_heading(heading, level=1)
        doc.add_paragraph(text)
        if heading == "Results":
            caption = doc.add_paragraph("Table 1. Synthetic estimates.")
            caption.style = styles["Caption"]
            table = doc.add_table(rows=len(TABLE), cols=len(TABLE[0]))
            table.style = "Table Grid"
            table.autofit = False
            widths = (Inches(1.5), Inches(1.2), Inches(2.1))
            for r, values in enumerate(TABLE):
                for c, value in enumerate(values):
                    cell = table.cell(r, c)
                    cell.width = widths[c]
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    cell.text = value
                    set_cell_margins(cell)
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if c < 2 else WD_ALIGN_PARAGRAPH.LEFT
                    if r == 0:
                        for run in cell.paragraphs[0].runs:
                            run.bold = True
            picture = doc.add_picture(str(figure), width=Inches(5.7))
            picture._inline.docPr.set("descr", "Bar chart of synthetic estimates for Groups A and B")
            fig_caption = doc.add_paragraph("Figure 1. Synthetic effect estimates for format verification.")
            fig_caption.style = styles["Caption"]
    doc.save(path)


def renderer_path() -> Path | None:
    root = Path.home() / ".codex/plugins/cache/openai-primary-runtime/documents"
    candidates = sorted(root.glob("*/skills/documents/render_docx.py"), reverse=True)
    return candidates[0] if candidates else None


def render_docx_pages(docx_path: Path, output: Path) -> tuple[list[Path], str]:
    output.mkdir(parents=True, exist_ok=True)
    if sys.platform == "win32":
        soffice = shutil.which("soffice")
        require(soffice is not None, "docx_visual_qa_pending: LibreOffice unavailable")
        from pdf2image import convert_from_path

        with tempfile.TemporaryDirectory(prefix="docx-profile-") as profile_dir:
            with tempfile.TemporaryDirectory(prefix="docx-pdf-") as pdf_dir:
                profile_uri = Path(profile_dir).resolve().as_uri()
                result = subprocess.run(
                    [
                        soffice,
                        f"-env:UserInstallation={profile_uri}",
                        "--headless",
                        "--norestore",
                        "--convert-to",
                        "pdf",
                        "--outdir",
                        pdf_dir,
                        str(docx_path),
                    ],
                    check=False,
                    capture_output=True,
                    text=True,
                    timeout=90,
                )
                pdf_path = Path(pdf_dir) / f"{docx_path.stem}.pdf"
                require(
                    result.returncode == 0 and pdf_path.is_file(),
                    "docx_visual_qa_pending: LibreOffice conversion failed",
                )
                pages = convert_from_path(str(pdf_path), dpi=180)
                paths = []
                for number, page in enumerate(pages, start=1):
                    path = output / f"page-{number}.png"
                    page.save(path)
                    paths.append(path)
        return paths, "libreoffice_pdf2image"

    renderer = renderer_path()
    require(renderer is not None, "docx_visual_qa_pending: documents renderer unavailable")
    try:
        subprocess.run(
            [sys.executable, str(renderer), str(docx_path), "--output_dir", str(output)],
            check=True,
            timeout=90,
        )
    except (subprocess.CalledProcessError, subprocess.TimeoutExpired) as error:
        raise AssertionError("docx_visual_qa_pending: documents renderer failed") from error
    return sorted(output.glob("page-*.png")), "documents_render_docx"


def structural_checks(markdown: Path, docx_path: Path, table_source: Path, figure: Path, manifest_path: Path) -> int:
    document = Document(docx_path)
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)
    for heading, content in SECTIONS.items():
        require(heading in paragraphs and content in text, f"DOCX section parity: {heading}")
    require("Table 1. Synthetic estimates." in paragraphs, "table caption")
    require("Figure 1. Synthetic effect estimates for format verification." in paragraphs, "figure caption")
    require([[cell.text for cell in row.cells] for row in document.tables[0].rows] == TABLE, "native table parity")
    with zipfile.ZipFile(docx_path) as package:
        document_xml = package.read("word/document.xml").decode("utf-8")
        require("<w:tbl>" in document_xml, "native Word table")
        require("<a:blip" in document_xml and "descr=\"Bar chart" in document_xml, "embedded figure and alt text")
        require(any(name.startswith("word/media/") for name in package.namelist()), "figure media part")
    manifest = {
        "schema_version": "research-article.v6",
        "canonical_markdown_ref": markdown.name,
        "canonical_content_digest": sha256(markdown),
        "docx_ref": docx_path.name,
        "docx_content_digest": sha256(docx_path),
        "docx_sync_status": "synchronized",
        "render_qa_status": "not_generated",
        "display_items": [
            {"display_id": "Table 1", "type": "table", "source_path": table_source.name, "source_digest": sha256(table_source), "placement": "main", "caption": "Synthetic estimates.", "callout": "Table 1", "availability": "available", "docx_embedding_status": "embedded"},
            {"display_id": "Figure 1", "type": "figure", "source_path": figure.name, "source_digest": sha256(figure), "placement": "main", "caption": "Synthetic effect estimates for format verification.", "callout": "Figure 1", "availability": "available", "alt_text": "Bar chart of synthetic estimates for Groups A and B", "docx_embedding_status": "embedded"},
        ],
    }
    manifest_path.write_text(yaml.safe_dump(manifest, sort_keys=False), encoding="utf-8", newline="\n")
    require("Table 1" in read_text(markdown) and "Figure 1" in read_text(markdown), "Markdown callouts")
    return 10


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8-sig")


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--render", action="store_true")
    parser.add_argument("--keep-dir", type=Path)
    args = parser.parse_args()

    require(CONTRACT.is_file(), "DOCX delivery contract exists")
    for marker in ("native Word tables", "render", "docx_visual_qa_pending", "content_drift"):
        require(marker in read_text(CONTRACT), f"contract marker: {marker}")

    temp_owner = None
    if args.keep_dir:
        work = args.keep_dir.resolve()
        tests_root = (REPO / "tests").resolve()
        require(work == tests_root or tests_root in work.parents, "--keep-dir must stay under tests/")
        require(not work.exists() or not any(work.iterdir()), "--keep-dir must be absent or empty")
        work.mkdir(parents=True, exist_ok=True)
    else:
        temp_owner = tempfile.TemporaryDirectory()
        work = Path(temp_owner.name)

    markdown = work / "manuscript-v001.md"
    table_source = work / "table-1.csv"
    figure = work / "figure-1.png"
    docx_path = work / "manuscript-v001.docx"
    manifest = work / "display-asset-manifest.yaml"
    make_markdown(markdown)
    make_table_source(table_source)
    make_figure(figure)
    build_docx(docx_path, figure)
    guards = structural_checks(markdown, docx_path, table_source, figure, manifest)

    render_pages: list[str] = []
    render_route = "not_requested"
    if args.render:
        output = work / "rendered"
        pages, render_route = render_docx_pages(docx_path, output)
        require(bool(pages), "rendered page PNGs")
        for page in pages:
            with Image.open(page) as image:
                require(image.width > 500 and image.height > 500, f"render dimensions: {page.name}")
            render_pages.append(str(page))
        data = yaml.safe_load(read_text(manifest))
        data["render_qa_status"] = "rendered_pending_visual_inspection"
        for item in data["display_items"]:
            item["render_status"] = "rendered_pending_visual_inspection"
        manifest.write_text(yaml.safe_dump(data, sort_keys=False), encoding="utf-8", newline="\n")
        guards += len(pages) + 1

    print(json.dumps({"guards": guards, "work_dir": str(work), "docx": str(docx_path), "render_route": render_route, "render_pages": render_pages}, ensure_ascii=False))
    if temp_owner is not None:
        temp_owner.cleanup()
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
